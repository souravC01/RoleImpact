from __future__ import annotations

import io
import math
import os
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE = Path(r"C:\Users\soura\.codex\.chatgpt-projects\g-p-6a7be822d72c81919ff60a55d05d082e")
QA = WORKSPACE / "qa" / "roleimpact_design_docs"
DELIVERABLES = WORKSPACE / "deliverables"
SCREENSHOTS = WORKSPACE / "qa" / "roleimpact_wireframes" / "screens"
WIREFRAMES_DOCX = DELIVERABLES / "RoleImpact_Low_Fidelity_Wireframes_v1.0.docx"
TECH_DOCX = DELIVERABLES / "RoleImpact_Technical_Design_v1.0.docx"

INK = "132438"
MUTED = "5F6B78"
BLUE = "315F82"
DARK_BLUE = "234963"
LIGHT_BLUE = "E9F0F5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D6DDE4"
DANGER = "963C43"
LIGHT_DANGER = "F8EDEF"
WARN = "806225"
LIGHT_WARN = "FAF4E5"
SAFE = "3F6D57"
LIGHT_SAFE = "EDF6F1"
WHITE = "FFFFFF"
BLACK = "000000"

CONTENT_WIDTH_DXA = 9360


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run(run, *, font="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic
    return run


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        data = kwargs.get(edge)
        if not data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in data:
                element.set(qn(f"w:{key}"), str(data[key]))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    cell.width = Inches(width_dxa / 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int], *, indent_dxa=120):
    assert sum(widths_dxa) == CONTENT_WIDTH_DXA, widths_dxa
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, *, header_fill=LIGHT_GRAY, font_size=9.2, header_color=INK):
    border = {"val": "single", "sz": "5", "space": "0", "color": MID_GRAY}
    for r_index, row in enumerate(table.rows):
        keep_row_together(row)
        if r_index == 0:
            repeat_header(row)
        for cell in row.cells:
            set_cell_border(cell, top=border, start=border, bottom=border, end=border)
            if r_index == 0:
                shade_cell(cell, header_fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.06
                for run in p.runs:
                    set_run(run, size=font_size, color=header_color if r_index == 0 else INK, bold=(r_index == 0))


def paragraph_border_bottom(paragraph, color=BLUE, size="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "7")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_image_alt(shape, title: str, description: str):
    doc_pr = shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_field(paragraph, code: str):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run(run, size=8.5, color=MUTED)


def configure_document(doc: Document, title: str, short_title: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Title", 26, INK, 0, 6),
        ("Subtitle", 13, MUTED, 0, 14),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # Real list definitions; no fake bullet characters in the document body.
    list_bullet = styles["List Bullet"]
    list_bullet.font.name = "Calibri"
    list_bullet.font.size = Pt(10.5)
    list_bullet.paragraph_format.left_indent = Inches(0.5)
    list_bullet.paragraph_format.first_line_indent = Inches(-0.25)
    list_bullet.paragraph_format.space_after = Pt(4)
    list_bullet.paragraph_format.line_spacing = 1.167
    list_number = styles["List Number"]
    list_number.font.name = "Calibri"
    list_number.font.size = Pt(10.5)
    list_number.paragraph_format.left_indent = Inches(0.5)
    list_number.paragraph_format.first_line_indent = Inches(-0.25)
    list_number.paragraph_format.space_after = Pt(4)
    list_number.paragraph_format.line_spacing = 1.167

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(short_title.upper()), size=8.5, color=MUTED, bold=True)
    paragraph_border_bottom(p, color=MID_GRAY, size="8")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("ROLEIMPACT  •  "), size=8.5, color=MUTED, bold=True)
    add_field(p, "PAGE")

    props = doc.core_properties
    props.title = title
    props.subject = "RoleImpact product design"
    props.author = "RoleImpact project"
    props.keywords = "RoleImpact, wireframes, technical design, access impact simulator"


def add_masthead(doc: Document, kicker: str, title: str, subtitle: str, meta: Sequence[tuple[str, str]]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(kicker.upper()), size=9, color=BLUE, bold=True)
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(title), size=26, color=INK, bold=True)
    p = doc.add_paragraph(style="Subtitle")
    set_run(p.add_run(subtitle), size=13, color=MUTED)
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(f"{label}: "), size=10, color=INK, bold=True)
        set_run(p.add_run(value), size=10, color=INK)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(12)
    paragraph_border_bottom(rule, color=BLUE, size="20")


def add_callout(doc: Document, label: str, text: str, *, fill=LIGHT_BLUE, accent=BLUE, keep=True):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, start={"val": "single", "sz": "26", "color": accent})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = keep
    set_run(p.add_run(f"{label}: "), size=10.5, color=accent, bold=True)
    set_run(p.add_run(text), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullet(doc: Document, text: str, *, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.keep_together = True
    set_run(p.add_run(text), size=10.2, color=INK)
    return p


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.keep_together = True
    set_run(p.add_run(text), size=10.2, color=INK)
    return p


def set_repeat_table_header(row):
    repeat_header(row)


def add_table_from_rows(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int], *, font_size=9.0, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    set_table_geometry(table, widths)
    style_table(table, header_fill=header_fill, font_size=font_size)
    return table


def add_screen_annotation_table(doc: Document, rows: Sequence[tuple[str, str]]):
    table = doc.add_table(rows=0, cols=2)
    for label, text in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = text
        shade_cell(cells[0], LIGHT_BLUE)
        for run in cells[0].paragraphs[0].runs:
            set_run(run, size=9.2, color=DARK_BLUE, bold=True)
        for run in cells[1].paragraphs[0].runs:
            set_run(run, size=9.2, color=INK)
    set_table_geometry(table, [1800, 7560])
    border = {"val": "single", "sz": "5", "space": "0", "color": MID_GRAY}
    for row in table.rows:
        keep_row_together(row)
        for cell in row.cells:
            set_cell_border(cell, top=border, start=border, bottom=border, end=border)
    return table


def crop_for_doc(source: Path, destination: Path, *, top=0, bottom=None):
    image = Image.open(source).convert("RGB")
    width, height = image.size
    bottom = height if bottom is None else min(bottom, height)
    cropped = image.crop((0, top, width, bottom))
    cropped.save(destination, quality=92)
    return destination


def build_wireframes_doc():
    QA.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc, "RoleImpact Low-Fidelity Wireframes v1.0", "RoleImpact — Low-Fidelity Wireframes")
    add_masthead(
        doc,
        "Product design specification",
        "RoleImpact Low-Fidelity Wireframes",
        "Six-screen interaction model for the Access Change Impact Simulator",
        (("Version", "v1.0"), ("Date", "August 12, 2026"), ("Status", "Approved PRD converted into build-ready screen structure")),
    )
    add_callout(doc, "Design intent", "Business consequences lead; relationship graphs support the explanation. Every destructive-looking action is explicitly framed as a read-only simulation against Harborline Commerce baseline v1.")
    doc.add_heading("1. Experience model", level=1)
    p = doc.add_paragraph()
    set_run(p.add_run("Primary path: "), bold=True, size=10.5)
    set_run(p.add_run("Dashboard → New Simulation → Review Proposal → Run Analysis → Results → Test Recommendation → Save Scenario"), size=10.5)
    doc.add_heading("Interaction decisions locked for implementation", level=2)
    for item in (
        "Use one persistent product shell: Dashboard, Explorer, and Simulations. Entity details and results are contextual views within that shell.",
        "Keep New Simulation to one focused page with three visible steps and an explicit safe-simulation notice.",
        "Place Business Impact before Technical Impact on Results; the graph follows the human-readable explanation path.",
        "Treat Test This Recommendation as part of the public MVP demo because it completes the approved user story and proves scenario isolation.",
        "Defer report export to the next milestone; keep Save Scenario in the MVP and preserve the report API boundary for later.",
    ):
        add_bullet(doc, item)
    doc.add_heading("Responsive behavior", level=2)
    add_bullet(doc, "Primary design target: 1280–1440 px laptop. At 736 px, the side navigation becomes a horizontal band and content stacks. At 360 px, all cards become single-column; the graph remains scrollable and the written conclusion remains complete.")
    add_bullet(doc, "Mobile is not a core MVP requirement, but narrow layouts must not clip actions or hide the textual graph equivalent.")

    screens = [
        ("1", "Dashboard", "dashboard.png", [
            ("Purpose", "Show organizational health, existing fragility, and recent scenario evidence before the user starts a simulation."),
            ("Primary action", "New Simulation."),
            ("Core data", "Entity counts, workflow statuses, single-person dependencies, recent simulations."),
            ("Key behavior", "Workflow and dependency rows open focused Explorer views; recent scenarios open saved Results."),
            ("States", "Loading skeletons, empty history, organization-load failure, no existing dependency risk."),
        ]),
        ("2", "Organization Explorer", "explorer.png", [
            ("Purpose", "Search entities and understand the relationship chain around an employee, role, app, permission, or workflow."),
            ("Primary action", "Open full details or begin a contextual simulation."),
            ("Core data", "Filterable entity list, affected-path graph, selected entity summary."),
            ("Key behavior", "Selecting a node updates the side panel; default graph is filtered to decision-relevant paths."),
            ("Accessibility", "The side panel and ordered relationship list provide the same information as the graph."),
        ]),
        ("3", "Entity Details", "entity.png", [
            ("Purpose", "Explain current effective access and business coverage before the user proposes a change."),
            ("Primary action", "Simulate a Change with the entity preselected."),
            ("Core data", "Status, team, region, shift, roles, effective permissions, applications, workflow steps, access origins."),
            ("Key behavior", "Tabs switch access perspectives; every effective permission can reveal its role-assignment origin."),
            ("States", "Inactive employee, no effective access, conflicting data marked inconclusive."),
        ]),
        ("4", "New Simulation", "simulation.png", [
            ("Purpose", "Collect one supported change and let the reviewer confirm the exact current-state relationships before analysis."),
            ("Primary action", "Run Impact Analysis."),
            ("Change types", "Offboard employee; Revoke role from employee; Remove permission from role."),
            ("Validation", "Targets must exist in baseline v1; role must be assigned; permission must belong to the selected role; contradictory input returns a field error."),
            ("Safety cue", "Persistent statement that the baseline is immutable and no real access is modified."),
        ]),
        ("5", "Simulation Results", "results.png", [
            ("Purpose", "Support an approve/reject/revise decision with a plain-language verdict and traceable evidence."),
            ("Primary action", "Test This Recommendation; Save Scenario is secondary."),
            ("Hierarchy", "Verdict → business impact → explanation path → focused graph → technical impact → safe mitigation."),
            ("Graph semantics", "Unchanged relationships are neutral; removed relationships are dashed red; blocked/degraded conclusions are labeled in text."),
            ("No recommendation", "Show unmet constraints and state that no safe automatic recommendation exists."),
        ]),
        ("6", "Scenario History", "history.png", [
            ("Purpose", "Preserve audit evidence and reopen original or recommendation-test branches."),
            ("Primary action", "Open a saved result; New Simulation remains available."),
            ("Core data", "Name, change, target, verdict, branch count, run timestamp, optional reviewer notes."),
            ("Key behavior", "Reopen is P0; duplicate and compare can follow after the first complete vertical slice."),
            ("Empty state", "Explain what gets saved and offer New Simulation."),
        ]),
    ]
    for number, name, file_name, notes in screens:
        doc.add_page_break()
        doc.add_heading(f"{number}. {name}", level=1)
        image_path = SCREENSHOTS / file_name
        picture = doc.add_picture(str(image_path), width=Inches(6.42))
        set_image_alt(picture, f"RoleImpact {name} wireframe", notes[0][1])
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(3)
        caption.paragraph_format.space_after = Pt(9)
        set_run(caption.add_run(f"Wireframe {number} — {name} (low fidelity; layout and hierarchy only)"), size=8.8, color=MUTED, italic=True)
        add_screen_annotation_table(doc, notes)

    doc.add_page_break()
    doc.add_heading("7. Cross-screen states and acceptance notes", level=1)
    rows = [
        ("Loading", "Use local skeletons; retain screen heading and do not replace results with a full-page spinner."),
        ("Invalid data", "Return an inconclusive analysis with the invalid entity/rule IDs. Never infer missing organizational facts."),
        ("No material impact", "Show severity Low or None, preserve the technical diff, and state that workflows remain operational."),
        ("No recommendation", "List unmet prerequisites and explain why every candidate was excluded."),
        ("Save conflict", "Scenario result is immutable; name and reviewer notes may be updated with optimistic concurrency."),
        ("Graph fallback", "Every graph has a textual relationship path and impact list for keyboard and screen-reader use."),
    ]
    add_table_from_rows(doc, ["State", "Required behavior"], rows, [1900, 7460], font_size=9.2, header_fill=LIGHT_BLUE)
    doc.add_heading("Screen-level acceptance checklist", level=2)
    for item in (
        "Primary Priya Sharma scenario can be completed in under three minutes without documentation.",
        "Each screen has visible focus, keyboard-reachable controls, and a meaningful page heading.",
        "No simulation action changes baseline data; baseline version is visible at proposal and result time.",
        "Results remain understandable with the graph hidden.",
        "At common laptop widths, no key control or result requires horizontal page scrolling.",
    ):
        add_bullet(doc, item)

    doc.save(WIREFRAMES_DOCX)
    return WIREFRAMES_DOCX


def add_code_block(doc: Document, code: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    keep_row_together(table.rows[0])
    shade_cell(cell, "F4F6F8")
    border = {"val": "single", "sz": "5", "space": "0", "color": MID_GRAY}
    set_cell_border(cell, top=border, start=border, bottom=border, end=border)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(code.strip("\n").splitlines()):
        if index:
            p.add_run().add_break()
        set_run(p.add_run(line), font="Consolas", size=8.4, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)
    return table


def build_architecture_diagram(path: Path):
    width, height = 1500, 720
    image = Image.new("RGB", (width, height), "#" + WHITE)
    draw = ImageDraw.Draw(image)
    try:
        font_title = ImageFont.truetype("segoeuib.ttf", 34)
        font_head = ImageFont.truetype("segoeuib.ttf", 26)
        font_body = ImageFont.truetype("segoeui.ttf", 21)
        font_small = ImageFont.truetype("segoeui.ttf", 18)
    except OSError:
        font_title = font_head = font_body = font_small = ImageFont.load_default()

    def box(x, y, w, h, title, lines, fill, stroke="#" + BLUE):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=20, fill=fill, outline=stroke, width=3)
        draw.text((x + 22, y + 18), title, fill="#" + INK, font=font_head)
        yy = y + 62
        for line in lines:
            draw.text((x + 22, yy), line, fill="#" + MUTED, font=font_small)
            yy += 28

    def arrow(x1, y1, x2, y2, label=""):
        draw.line((x1, y1, x2, y2), fill="#" + BLUE, width=4)
        angle = math.atan2(y2 - y1, x2 - x1)
        size = 13
        p1 = (x2 - size * math.cos(angle - .55), y2 - size * math.sin(angle - .55))
        p2 = (x2 - size * math.cos(angle + .55), y2 - size * math.sin(angle + .55))
        draw.polygon([(x2, y2), p1, p2], fill="#" + BLUE)
        if label:
            tw = draw.textbbox((0, 0), label, font=font_small)[2]
            draw.rectangle(((x1 + x2 - tw) / 2 - 8, (y1 + y2) / 2 - 14, (x1 + x2 + tw) / 2 + 8, (y1 + y2) / 2 + 12), fill="#" + WHITE)
            draw.text(((x1 + x2 - tw) / 2, (y1 + y2) / 2 - 13), label, fill="#" + MUTED, font=font_small)

    draw.text((55, 35), "RoleImpact deployment and module boundaries", fill="#" + INK, font=font_title)
    box(55, 130, 330, 185, "React web", ["Vite + TypeScript", "TanStack Query", "React Flow graph"], "#" + LIGHT_BLUE)
    box(550, 130, 390, 185, "Spring Boot API", ["REST controllers", "Application services", "OpenAPI validation"], "#" + LIGHT_GRAY)
    box(1095, 130, 350, 185, "PostgreSQL", ["Baseline graph tables", "Saved simulations", "JSONB result snapshot"], "#" + LIGHT_SAFE, stroke="#" + SAFE)
    box(550, 430, 390, 190, "Impact engine", ["Pure deterministic module", "No HTTP / no database", "Input snapshot → result"], "#" + LIGHT_WARN, stroke="#" + WARN)
    box(1095, 430, 350, 190, "Report adapter (P1)", ["Consumes saved result", "No engine recalculation", "PDF / DOCX later"], "#" + LIGHT_DANGER, stroke="#" + DANGER)
    arrow(385, 222, 550, 222, "HTTPS / JSON")
    arrow(940, 222, 1095, 222, "JPA / SQL")
    arrow(745, 315, 745, 430, "typed snapshot")
    arrow(940, 525, 1095, 525, "saved result")
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path)
    return path


def build_schema_diagram(path: Path):
    width, height = 1700, 1050
    image = Image.new("RGB", (width, height), "#" + WHITE)
    draw = ImageDraw.Draw(image)
    try:
        font_title = ImageFont.truetype("segoeuib.ttf", 34)
        font_head = ImageFont.truetype("segoeuib.ttf", 22)
        font_body = ImageFont.truetype("consola.ttf", 17)
    except OSError:
        font_title = font_head = font_body = ImageFont.load_default()

    entities = {
        "organizations": (40, 115, ["id PK", "slug UQ", "current_version"]),
        "teams": (40, 350, ["id PK", "organization_id FK", "department", "manager_employee_id FK"]),
        "employees": (390, 310, ["id PK", "team_id FK", "status", "region", "shift"]),
        "employee_roles": (765, 160, ["employee_id PK/FK", "role_id PK/FK", "assigned_at"]),
        "roles": (1125, 120, ["id PK", "organization_id FK", "sensitivity", "owner_employee_id FK"]),
        "role_permissions": (1125, 355, ["role_id PK/FK", "permission_id PK/FK"]),
        "permissions": (1465, 310, ["id PK", "application_id FK", "resource_id FK", "action", "sensitivity"]),
        "applications": (1465, 575, ["id PK", "organization_id FK", "name", "category"]),
        "capabilities": (805, 500, ["id PK", "organization_id FK", "name"]),
        "capability_permissions": (1115, 590, ["capability_id PK/FK", "permission_id PK/FK"]),
        "workflows": (390, 665, ["id PK", "organization_id FK", "criticality", "owner_employee_id FK"]),
        "workflow_steps": (750, 750, ["id PK", "workflow_id FK", "required_capability_id FK", "min_actors / resilience_target", "required_region/shift/app"]),
        "workflow_constraints": (1115, 825, ["id PK", "workflow_id FK", "type", "parameters JSONB"]),
        "simulations": (40, 800, ["id PK", "organization_id FK", "baseline_version", "parent_simulation_id FK", "change JSONB", "result JSONB"]),
    }
    sizes = {}
    for name, (x, y, fields) in entities.items():
        w = 300 if name not in ("workflow_steps", "simulations") else 330
        h = 62 + 29 * len(fields)
        sizes[name] = (x, y, w, h)

    def center(name, side):
        x, y, w, h = sizes[name]
        return {"left": (x, y + h / 2), "right": (x + w, y + h / 2), "top": (x + w / 2, y), "bottom": (x + w / 2, y + h)}[side]

    relations = [
        ("organizations", "right", "employee_roles", "left"),
        ("organizations", "bottom", "teams", "top"),
        ("teams", "right", "employees", "left"),
        ("employees", "right", "employee_roles", "left"),
        ("employee_roles", "right", "roles", "left"),
        ("roles", "bottom", "role_permissions", "top"),
        ("role_permissions", "right", "permissions", "left"),
        ("permissions", "bottom", "applications", "top"),
        ("permissions", "left", "capability_permissions", "right"),
        ("capability_permissions", "left", "capabilities", "right"),
        ("capabilities", "bottom", "workflow_steps", "top"),
        ("workflows", "right", "workflow_steps", "left"),
        ("workflow_steps", "right", "workflow_constraints", "left"),
        ("organizations", "bottom", "simulations", "top"),
    ]
    draw.text((40, 35), "Principal MVP relational paths (effective access is derived)", fill="#" + INK, font=font_title)
    for source, ss, target, ts in relations:
        x1, y1 = center(source, ss)
        x2, y2 = center(target, ts)
        mid_x = (x1 + x2) / 2
        draw.line((x1, y1, mid_x, y1, mid_x, y2, x2, y2), fill="#" + MID_GRAY, width=3)
    for name, (x, y, fields) in entities.items():
        _, _, w, h = sizes[name]
        fill = LIGHT_BLUE if name in ("employees", "roles", "permissions", "capabilities", "workflows", "workflow_steps") else (LIGHT_WARN if name == "simulations" else LIGHT_GRAY)
        stroke = WARN if name == "simulations" else BLUE
        draw.rounded_rectangle((x, y, x + w, y + h), radius=14, fill="#" + fill, outline="#" + stroke, width=3)
        draw.rectangle((x, y, x + w, y + 46), fill="#" + stroke)
        draw.text((x + 14, y + 10), name, fill="#" + WHITE, font=font_head)
        yy = y + 56
        for field in fields:
            draw.text((x + 14, yy), field, fill="#" + INK, font=font_body)
            yy += 28
    image.save(path)
    return path


def build_technical_doc():
    QA.mkdir(parents=True, exist_ok=True)
    architecture = build_architecture_diagram(QA / "architecture.png")
    schema = build_schema_diagram(QA / "schema.png")
    doc = Document()
    configure_document(doc, "RoleImpact Technical Design v1.0", "RoleImpact — Technical Design")
    add_masthead(
        doc,
        "Technical design specification",
        "RoleImpact Technical Design",
        "Final data schema, deterministic impact engine, REST boundaries, tests, and implementation stack",
        (("Version", "v1.0"), ("Date", "August 12, 2026"), ("Architecture", "Two deployables + one PostgreSQL database; engine is a pure backend module")),
    )
    add_callout(doc, "Recommendation", "Build RoleImpact as a focused Java 21 / Spring Boot modular monolith with a React + TypeScript client and PostgreSQL. Keep the graph in relational tables, derive effective access, and avoid Neo4j, Kafka, Redis, Kubernetes, and real IAM integration for the MVP.")
    doc.add_heading("1. Decisions at a glance", level=1)
    decisions = [
        ("Backend", "Java 21 + Spring Boot 3.5.x; Spring MVC, Bean Validation, Spring Data JPA, Flyway."),
        ("Frontend", "React + TypeScript + Vite; TanStack Query, React Hook Form + Zod, React Flow, CSS/Tailwind utility layer as preferred."),
        ("Database", "PostgreSQL 17; normalized baseline graph plus JSONB change/result snapshots for saved simulations."),
        ("Architecture", "One Spring Boot API, one static React app, one database. Domain modules inside the API; no microservices."),
        ("Engine", "Pure deterministic Java module: OrganizationSnapshot + SimulationChange → ImpactResult. No repository or HTTP dependency."),
        ("Deployment", "Docker Compose locally; Render static site + Docker web service + managed Postgres for the public portfolio demo."),
        ("Scope", "Recommendation testing is in the public MVP. Report export is P1 and consumes saved results without recalculation."),
    ]
    add_table_from_rows(doc, ["Decision", "Final choice"], decisions, [1800, 7560], font_size=9.1, header_fill=LIGHT_BLUE)

    doc.add_heading("2. Architecture and boundaries", level=1)
    p = doc.add_paragraph()
    set_run(p.add_run("Architecture style. "), bold=True)
    set_run(p.add_run("A small two-deployable system keeps the product understandable while adding meaningful Java/API depth for the target software-engineering roles. The backend remains a modular monolith; modules communicate through typed Java interfaces, not network calls."))
    picture = doc.add_picture(str(architecture), width=Inches(6.5))
    set_image_alt(picture, "RoleImpact architecture diagram", "React and TypeScript web client calls a Spring Boot API. The API reads PostgreSQL and passes an immutable snapshot to a pure deterministic impact engine. A P1 report adapter consumes only saved results.")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(caption.add_run("Figure 1 — Runtime and module boundaries"), size=8.8, color=MUTED, italic=True)
    doc.add_heading("Backend modules", level=2)
    module_rows = [
        ("catalog", "Read/query organization entities and assemble immutable OrganizationSnapshot objects."),
        ("simulation", "Validate a proposed change, invoke the engine, persist immutable result snapshots, and manage branches/notes."),
        ("impactengine", "Apply the change, derive effective access, evaluate workflow coverage, create explanations, graph diffs, and recommendations."),
        ("reporting (P1)", "Render an already-saved result. It must not invoke the engine or mutate scenario data."),
        ("shared", "IDs, enums, error envelope, clock abstraction, canonical JSON/hash utilities."),
    ]
    add_table_from_rows(doc, ["Module", "Responsibility"], module_rows, [1900, 7460], font_size=9.2)
    doc.add_heading("Boundary rules", level=2)
    for item in (
        "The React client calls only /api/v1 and never reads PostgreSQL directly.",
        "Controllers map HTTP DTOs to application commands. They do not contain workflow rules or repository queries.",
        "The impact engine accepts a complete in-memory snapshot and returns a complete result. It imports no JPA, Spring MVC, or database classes.",
        "Repositories return domain projections; JPA entities are not serialized to the client.",
        "Saved result JSON is versioned. Reopening a scenario reads the saved result; it does not silently recalculate against a newer baseline.",
        "Only name and reviewer notes are mutable after save. Inputs, result, baseline version, and branch ancestry are immutable.",
    ):
        add_bullet(doc, item)

    doc.add_page_break()
    doc.add_heading("3. Final data schema", level=1)
    add_callout(doc, "Schema rule", "Relational tables are the source of truth. Effective permissions, step eligibility, workflow status, explanation paths, and graph diffs are derived. They are stored only inside an immutable simulation result for reproducibility.")
    picture = doc.add_picture(str(schema), width=Inches(6.5))
    set_image_alt(picture, "RoleImpact principal schema paths", "Principal foreign-key paths connecting organizations, teams, employees, roles, permissions, capabilities, workflows, workflow steps, constraints, and immutable saved simulations. Complete table definitions follow the figure.")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(caption.add_run("Figure 2 — Principal foreign-key paths; complete table definitions follow"), size=8.8, color=MUTED, italic=True)

    doc.add_heading("3.1 Organization and access tables", level=2)
    access_rows = [
        ("organizations", "id uuid PK; slug varchar UQ; name; current_version int; created_at", "Root aggregate; MVP seed contains Harborline Commerce."),
        ("teams", "id uuid PK; organization_id FK; name; department; manager_employee_id FK nullable", "Unique (organization_id, name)."),
        ("employees", "id uuid PK; organization_id FK; team_id FK; employee_no UQ; name; email; status; region; shift", "status ACTIVE/INACTIVE; region and shift are constrained enums."),
        ("roles", "id uuid PK; organization_id FK; name; description; sensitivity; owner_employee_id FK nullable", "Unique role name per organization."),
        ("employee_roles", "employee_id FK; role_id FK; assigned_at; assigned_by", "Composite PK (employee_id, role_id). No overlapping history in MVP."),
        ("applications", "id uuid PK; organization_id FK; name; category; owner_employee_id FK nullable", "Unique app name per organization."),
        ("resources", "id uuid PK; application_id FK; name; resource_type", "Unique (application_id, name)."),
        ("permissions", "id uuid PK; organization_id FK; application_id FK; resource_id FK; action; sensitivity", "Unique (application_id, resource_id, action)."),
        ("role_permissions", "role_id FK; permission_id FK", "Composite PK. Role definition change removes one row in scenario memory only."),
    ]
    add_table_from_rows(doc, ["Table", "Important columns", "Constraint / purpose"], access_rows, [1650, 4420, 3290], font_size=8.25, header_fill=LIGHT_BLUE)

    doc.add_heading("3.2 Capabilities and workflow tables", level=2)
    workflow_rows = [
        ("capabilities", "id uuid PK; organization_id FK; name; description", "Business-facing ability such as Approve high-value payment."),
        ("capability_permissions", "capability_id FK; permission_id FK", "Composite PK. Any mapped effective permission satisfies the capability."),
        ("workflows", "id uuid PK; organization_id FK; name; criticality; owner_employee_id FK", "criticality CRITICAL/HIGH/MEDIUM/LOW."),
        ("workflow_steps", "id uuid PK; workflow_id FK; step_key; name; position; required_capability_id FK; minimum_actors; resilience_target; required_department/region/shift/application_id nullable", "Unique (workflow_id, step_key) and (workflow_id, position); resilience_target >= minimum_actors."),
        ("workflow_constraints", "id uuid PK; workflow_id FK; type; parameters jsonb", "Typed, validated parameters for SOD, EXCLUDE_EMPLOYEE, and DIFFERENT_ACTORS."),
        ("workflow_step_constraints", "workflow_step_id FK; constraint_id FK", "Composite PK; connects a constraint to participating steps."),
    ]
    add_table_from_rows(doc, ["Table", "Important columns", "Constraint / purpose"], workflow_rows, [2200, 4240, 2920], font_size=8.25, header_fill=LIGHT_BLUE)

    doc.add_heading("3.3 Simulation persistence", level=2)
    simulation_rows = [
        ("simulations", "id uuid PK; organization_id FK; parent_simulation_id FK nullable; baseline_version; engine_version; idempotency_key nullable; request_hash; change_type; change_payload jsonb; result_status; result_payload jsonb; severity; name; reviewer_notes; created_at; completed_at", "Original and tested-recommendation branches use the same table. Inputs and results are immutable; name/notes may change."),
        ("organization_versions", "organization_id FK; version int; content_hash; created_at", "Composite PK; records the seed version used for reproducibility without cloning every baseline row."),
    ]
    add_table_from_rows(doc, ["Table", "Important columns", "Constraint / purpose"], simulation_rows, [2200, 4740, 2420], font_size=8.15, header_fill=LIGHT_WARN)
    doc.add_heading("Indexes and integrity constraints", level=2)
    for item in (
        "Indexes on every foreign key plus employees(organization_id, status), employee_roles(role_id, employee_id), role_permissions(permission_id, role_id), workflow_steps(required_capability_id), and simulations(organization_id, created_at desc).",
        "CHECK constraints enforce known enums, positive positions, minimum_actors ≥ 1, and baseline_version ≥ 1.",
        "change_payload and result_payload require schema_version. Unique (organization_id, idempotency_key) plus request_hash makes retried POSTs safe while still allowing the same scenario to be intentionally run again under a new key.",
        "Foreign keys use RESTRICT for baseline catalog deletion. The MVP seeds data through Flyway and does not expose mutation endpoints.",
        "No graph database. The seed is small; indexed joins plus in-memory traversal are simpler, portable, and easier to test.",
    ):
        add_bullet(doc, item)

    doc.add_heading("4. Impact-engine contract", level=1)
    p = doc.add_paragraph()
    set_run(p.add_run("Contract invariant. "), bold=True)
    set_run(p.add_run("For the same canonical baseline snapshot, engine version, and change, the engine returns byte-equivalent canonical JSON and the same result hash. Timestamps and generated IDs are supplied by the application layer, not created inside the engine."))
    doc.add_heading("4.1 Java boundary", level=2)
    add_code_block(doc, """
public interface ImpactEngine {
  ImpactResult analyze(OrganizationSnapshot baseline, SimulationChange change);
}

public sealed interface SimulationChange permits
    OffboardEmployee, RevokeEmployeeRole, RemoveRolePermission {}

record OffboardEmployee(UUID employeeId) implements SimulationChange {}
record RevokeEmployeeRole(UUID employeeId, UUID roleId) implements SimulationChange {}
record RemoveRolePermission(UUID roleId, UUID permissionId) implements SimulationChange {}
""")
    doc.add_heading("4.2 Input snapshot", level=2)
    input_rows = [
        ("schemaVersion", "string", "Contract version, initially 1.0."),
        ("organizationId / baselineVersion", "UUID / integer", "Stable baseline identity."),
        ("employees / teams", "arrays", "Status and organizational attributes required for eligibility."),
        ("roles / assignments", "arrays", "Role definitions and current employee-role edges."),
        ("permissions / apps / resources", "arrays", "Technical access graph."),
        ("capabilities / mappings", "arrays", "Qualifying permission alternatives for each business ability."),
        ("workflows / steps / constraints", "arrays", "Ordered requirements, minimum coverage, and SOD rules."),
    ]
    add_table_from_rows(doc, ["Field", "Type", "Meaning"], input_rows, [2200, 1600, 5560], font_size=9.0)
    doc.add_heading("4.3 Output contract", level=2)
    output_rows = [
        ("resultStatus", "COMPLETE | INCONCLUSIVE", "INCONCLUSIVE includes typed data errors and no fabricated conclusion."),
        ("overallSeverity", "CRITICAL | HIGH | MEDIUM | LOW | NONE", "Computed from workflow status deltas and technical-only change."),
        ("executiveSummary", "structured counts + message key + parameters", "Frontend may render a sentence without changing the evidence."),
        ("changeSet", "removed/added nodes and edges", "Canonical direct changes applied to scenario memory."),
        ("technicalImpact", "roles, permissions, apps, resources, identities", "Before/after effective-access diff."),
        ("workflowImpacts", "workflow and step evaluations", "Baseline/scenario status, eligible actor IDs, failures, consequences."),
        ("explanationPaths", "ordered typed nodes/edges", "Every blocked/degraded conclusion has at least one exact path."),
        ("graphDiff", "focused node/edge DTOs", "Presentation projection; not the source of truth for the verdict."),
        ("recommendations", "0–2 candidates", "Actions, restored steps, required grants, exclusions, rank breakdown."),
        ("diagnostics", "engineVersion, resultHash, durationMs", "durationMs is metadata excluded from deterministic result hash."),
    ]
    add_table_from_rows(doc, ["Field", "Shape", "Guarantee"], output_rows, [1950, 2300, 5110], font_size=8.8, header_fill=LIGHT_BLUE)

    doc.add_heading("4.4 Evaluation pipeline", level=2)
    for step in (
        "Validate the snapshot and change. Return INCONCLUSIVE with entity/rule references for contradictions.",
        "Compute baseline effective permission sets and workflow-step eligibility.",
        "Copy the in-memory edge sets and apply exactly one proposed change.",
        "Find affected identities and capabilities from changed assignments or role-permission edges.",
        "Recompute only connected workflow steps, then evaluate minimum coverage and SOD constraints.",
        "Compare baseline and scenario to classify Operational, Degraded, or Blocked and calculate overall severity.",
        "Build explanation paths and a focused graph diff from stored predecessor links.",
        "Generate deterministic candidate actions; filter unsafe candidates, score the remainder, and return at most two.",
        "Canonicalize output ordering, compute resultHash, and return the immutable ImpactResult.",
    ):
        add_numbered(doc, step)

    doc.add_heading("4.5 Eligibility, status, and severity", level=2)
    rules = [
        ("Eligible actor", "Active employee AND qualifying effective permission AND department/region/shift prerequisites AND not excluded. For assignment-level SOD, the proposed actor combination must remain satisfiable."),
        ("Operational", "Every required step meets minimum_actors and configured resilience_target; all mandatory constraints are satisfiable."),
        ("Degraded", "Workflow remains executable, but at least one step falls below its resilience_target or becomes a single-person dependency."),
        ("Blocked", "At least one required step has no eligible actor or mandatory constraints cannot be satisfied."),
        ("Severity", "Critical workflow blocked → Critical; non-critical blocked or critical degraded → High; non-critical degraded/new sole dependency → Medium; access-only diff → Low; no effective change → None."),
    ]
    add_table_from_rows(doc, ["Rule", "Final definition"], rules, [1900, 7460], font_size=9.0, header_fill=LIGHT_GRAY)

    doc.add_heading("4.6 Recommendation ranking", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("Candidate action in MVP: "), bold=True)
    set_run(p.add_run("assign an existing role to one existing active employee. The engine does not create roles, clone permissions, or modify workflow policy."))
    rank_rows = [
        ("Safety gate", "Required", "Active; prerequisites satisfied; application prerequisites satisfied; no SOD conflict; candidate not excluded."),
        ("New permissions", "Primary", "Fewer permissions added ranks first."),
        ("Organizational fit", "Secondary", "Same department/team, then region/shift fit."),
        ("Existing app access", "Secondary", "Prefer a candidate already using the application."),
        ("Coverage restored", "Secondary", "More blocked/degraded steps restored ranks higher."),
        ("Resilience", "Tie-break", "Prefer improvement in backup coverage; final tie-break is employee UUID ascending."),
    ]
    add_table_from_rows(doc, ["Stage", "Priority", "Rule"], rank_rows, [1800, 1500, 6060], font_size=8.9, header_fill=LIGHT_WARN)
    add_callout(doc, "No-safe-option behavior", "Return recommendations=[] plus excludedCandidateReasons grouped by unmet constraint. The API and UI must say that no safe automatic recommendation exists.", fill=LIGHT_DANGER, accent=DANGER)

    doc.add_heading("5. REST API boundaries", level=1)
    add_callout(doc, "API style", "JSON over HTTPS under /api/v1. OpenAPI is generated from the Spring API and used to generate the frontend TypeScript client. No GraphQL is needed for the MVP.")
    endpoint_rows = [
        ("GET", "/api/v1/dashboard", "Counts, workflow health, dependencies, recent scenarios", "DashboardSummary"),
        ("GET", "/api/v1/entities", "type, q, team, status, page, size", "Paged<EntitySummary>"),
        ("GET", "/api/v1/entities/{type}/{id}", "Entity detail, access origins, workflow links", "EntityDetail"),
        ("GET", "/api/v1/graph", "focusType, focusId, scope=affected|neighborhood", "GraphProjection"),
        ("POST", "/api/v1/simulations/validate", "Change DTO + baselineVersion", "ProposalPreview; no persistence"),
        ("POST", "/api/v1/simulations", "Change DTO, baselineVersion, Idempotency-Key", "201 SimulationResultResource"),
        ("GET", "/api/v1/simulations", "filters, cursor/page, sort", "Paged<SimulationSummary>"),
        ("GET", "/api/v1/simulations/{id}", "Saved immutable result and branch summary", "SimulationResultResource"),
        ("PATCH", "/api/v1/simulations/{id}", "name, reviewerNotes, expectedUpdatedAt", "SimulationMetadata"),
        ("POST", "/api/v1/simulations/{id}/branches", "Recommendation ID or explicit candidate action", "201 branched SimulationResultResource"),
        ("GET", "/api/v1/simulations/{id}/report", "P1; format=pdf|docx", "Generated report stream"),
    ]
    add_table_from_rows(doc, ["Verb", "Path", "Input / responsibility", "Response"], endpoint_rows, [900, 3200, 3360, 1900], font_size=7.95, header_fill=LIGHT_BLUE)

    doc.add_heading("5.1 Change request", level=2)
    add_code_block(doc, """
{
  "schemaVersion": "1.0",
  "organizationId": "harborline-id",
  "baselineVersion": 1,
  "change": {
    "type": "REVOKE_EMPLOYEE_ROLE",
    "employeeId": "priya-id",
    "roleId": "finance-approver-id"
  }
}
""")
    doc.add_heading("5.2 HTTP behavior", level=2)
    http_rows = [
        ("200/201", "Successful query or newly computed saved simulation/branch."),
        ("400", "Malformed JSON or invalid enum/field shape."),
        ("404", "Organization, entity, simulation, or recommendation ID does not exist."),
        ("409", "Baseline version is stale; idempotency key reused with different input; metadata optimistic-lock conflict."),
        ("422", "Well-formed but invalid scenario: role not assigned, permission not on role, or contradictory baseline reference."),
        ("500", "Unexpected fault with trace ID; no internal stack or sensitive data returned."),
    ]
    add_table_from_rows(doc, ["Status", "Meaning"], http_rows, [1500, 7860], font_size=9.2)
    p = doc.add_paragraph()
    set_run(p.add_run("Error envelope: "), bold=True)
    set_run(p.add_run("{ code, message, traceId, fieldErrors[], entityRefs[] }. All timestamps are UTC ISO-8601; all IDs are UUID strings; unknown response fields must be ignored by clients."))

    doc.add_heading("6. Test strategy", level=1)
    p = doc.add_paragraph()
    set_run(p.add_run("Testing principle. "), bold=True)
    set_run(p.add_run("Correctness evidence matters more than coverage percentage. The suite must prove isolation, deterministic output, complete explanations, and the three supported change types."))
    test_rows = [
        ("Pure unit", "JUnit 5 + AssertJ", "Effective access, eligibility predicates, SOD satisfiability, status/severity, graph diff, explanation paths, recommendation filtering/ranking.", "Every rule branch and authored scenario."),
        ("Property/invariant", "jqwik or custom generated fixtures", "Same input → same hash; baseline unchanged; recommendation never worsens severity; every degraded/blocked result has a path; no duplicate nodes/edges.", "200+ generated cases in CI."),
        ("Persistence integration", "Spring Boot + Testcontainers PostgreSQL", "Flyway migrations, repository projections, JSONB round-trip, transaction rollback, idempotency, branch ancestry.", "Real Postgres, never H2."),
        ("API contract", "MockMvc/RestAssured + OpenAPI validation", "DTO validation, status codes, error envelope, pagination, stale baseline conflict.", "All public endpoints."),
        ("Frontend unit", "Vitest + Testing Library", "Form conditional fields, status rendering, explanation list, graph text alternative, recommendation states.", "Critical components and hooks."),
        ("End-to-end", "Playwright", "Primary Priya flow, all three change types, no-impact case, no-safe-recommendation, save/reopen branch.", "Chromium on every PR; broader matrix nightly if desired."),
        ("Accessibility", "axe-core + keyboard scripts", "Focus order, labels, contrast, dialogs, table semantics, graph alternative.", "No serious/critical issues on six screens."),
        ("Performance", "JUnit benchmark harness + timed API smoke", "Seeded simulation p95 under 2 s; dashboard p95 under 500 ms after warm-up.", "Stable local/CI profile; report hardware."),
    ]
    add_table_from_rows(doc, ["Layer", "Tool", "What it proves", "Gate"], test_rows, [1200, 1800, 4350, 2010], font_size=7.8, header_fill=LIGHT_BLUE)

    doc.add_heading("6.1 Golden acceptance scenarios", level=2)
    scenario_rows = [
        ("Revoke Finance Approver from Priya", "Vendor Payment blocked in evening; Month-End Close degraded; Critical; complete path; Bob recommended only if eligible."),
        ("Offboard sole release approver", "Production Deployment blocked or degraded according to remaining actors and criticality."),
        ("Remove refund.approve from Senior Support", "Every role holder recalculated; high-value refund step loses coverage; impacted identity list complete."),
        ("Revoke redundant viewer role", "Technical diff exists; workflows remain operational; severity Low."),
        ("No safe candidate", "recommendations empty; unmet shift/SOD/application constraints displayed."),
        ("Test Bob branch", "Parent unchanged; branch restores coverage; branch references parent; result can be reopened exactly."),
    ]
    add_table_from_rows(doc, ["Scenario", "Required assertion"], scenario_rows, [2900, 6460], font_size=8.8, header_fill=LIGHT_WARN)

    doc.add_heading("6.2 CI quality gates", level=2)
    for item in (
        "Backend unit and integration tests pass; Flyway migrates a blank database and seed dataset.",
        "Frontend typecheck, lint, unit tests, and production build pass.",
        "Generated OpenAPI client is current; a CI diff fails when backend contract changes are not regenerated.",
        "Playwright primary scenario passes against the full Docker Compose stack.",
        "No high-severity dependency or secret-scan findings; container health checks pass.",
        "The benchmark records p50/p95 and seeded dataset size; do not claim a performance number before measuring it.",
    ):
        add_bullet(doc, item)

    doc.add_page_break()
    doc.add_heading("7. Implementation stack", level=1)
    stack_rows = [
        ("Language/runtime", "Java 21 LTS", "Relevant to banking/platform roles and sufficient for modern Spring; avoids adopting a newer Java only for novelty."),
        ("API framework", "Spring Boot 3.5.x + Maven", "Mature stable line compatible with Java 21; Spring MVC, Validation, Actuator, Data JPA."),
        ("Database", "PostgreSQL 17 + Flyway", "Supported major version; SQL joins and JSONB snapshots cover the MVP without a graph database."),
        ("Frontend", "React + TypeScript + Vite", "Focused SPA, fast local setup, clear separation from Java API."),
        ("Server state/forms", "TanStack Query; React Hook Form + Zod", "Explicit request caching, mutations, conditional validation, and typed form errors."),
        ("Graph", "@xyflow/react", "Focused affected-path graph with keyboard-aware nodes; written explanation remains authoritative."),
        ("UI", "Tailwind CSS + small reusable components", "Fast portfolio polish; avoid bringing a complete enterprise design system."),
        ("Tests", "JUnit 5, AssertJ, Testcontainers, Vitest, Testing Library, Playwright, axe-core", "Covers domain, real Postgres, API, UI, E2E, and accessibility."),
        ("Local/CI", "Docker Compose + GitHub Actions", "One-command seed and repeatable full-stack checks."),
        ("Deploy", "Render static site + Docker web service + managed Postgres", "One platform for a public demo; use a persistent paid DB or reseed an ephemeral demo DB."),
        ("Observability", "Spring Boot Actuator + structured JSON logs + traceId", "Enough operational evidence without adding a metrics platform in MVP."),
    ]
    add_table_from_rows(doc, ["Layer", "Choice", "Rationale"], stack_rows, [1700, 2800, 4860], font_size=8.55, header_fill=LIGHT_BLUE)
    add_callout(doc, "Version policy", "Pin exact dependency versions in lockfiles/build files when implementation starts. This design names stable lines, not an instruction to upgrade automatically during the build.", fill=LIGHT_WARN, accent=WARN)

    doc.add_heading("8. Repository and delivery plan", level=1)
    add_code_block(doc, """
roleimpact/
├── frontend/              # React + TypeScript SPA
├── backend/               # Spring Boot modular monolith
│   └── src/main/java/.../{catalog,simulation,impactengine,reporting,shared}
├── docs/                  # PRD, wireframes, technical design, API examples
├── infra/                 # docker-compose.yml and Render blueprint
├── .github/workflows/     # backend, frontend, e2e, security checks
└── README.md              # problem, demo, architecture, metrics, trade-offs
""")
    phases = [
        ("1. Foundation", "Create repo, Flyway schema/seed, snapshot assembler, OpenAPI shell, frontend app shell.", "Dashboard reads seeded Harborline data."),
        ("2. Vertical slice", "Revoke-role change, effective access, Vendor Payment rules, basic Results.", "Priya scenario returns correct blocked path."),
        ("3. Engine completeness", "Other change types, all workflows, constraints, severity, explanations, recommendations.", "Golden engine suite passes."),
        ("4. Product flow", "Explorer, entity detail, result graph, branch test, save/reopen, error/empty states.", "Six screens work end to end."),
        ("5. Portfolio proof", "E2E/a11y/performance, Docker, deployment, README, demo video, measured resume bullets.", "Public demo and CI are reproducible."),
    ]
    add_table_from_rows(doc, ["Phase", "Build", "Exit condition"], phases, [1800, 4760, 2800], font_size=8.7, header_fill=LIGHT_GRAY)

    doc.add_heading("9. Explicitly deferred", level=1)
    for item in (
        "Real identity-provider, HR, ERP, cloud, or ticketing integrations.",
        "Authentication and multi-tenant administration; the public demo is read-only and synthetic.",
        "Neo4j, Kafka, Redis, Kubernetes, background queues, or multiple backend services.",
        "Generic policy language, time-bound grants, complex ABAC, and arbitrary workflow authoring.",
        "Report export until the core web experience and saved-result contract are complete.",
        "AI-generated recommendations. Any future language summary must cite deterministic evidence and remain non-authoritative.",
    ):
        add_bullet(doc, item)

    doc.add_heading("10. Definition of ready for implementation", level=1)
    for item in (
        "The six wireframes and technical boundaries in this package are the implementation baseline.",
        "Schema changes require a Flyway migration and an updated diagram/table in this document.",
        "Engine contract changes require golden-test updates and OpenAPI response review.",
        "A vertical slice starts with Priya's role revocation; secondary screens must not delay that path.",
        "A stretch feature enters only after the end-to-end primary scenario, tests, and deployment work.",
    ):
        add_bullet(doc, item)

    doc.add_heading("Appendix A. Authoritative implementation references", level=1)
    refs = [
        "Spring Boot stable versions and Java requirements: https://docs.spring.io/spring-boot/system-requirements.html",
        "PostgreSQL supported versions: https://www.postgresql.org/support/versioning/",
        "Vite React/TypeScript templates: https://vite.dev/guide/",
        "React Flow quick start and @xyflow/react package: https://reactflow.dev/learn",
        "Playwright Test documentation: https://playwright.dev/docs/api/class-test",
        "Testcontainers Spring Boot/Postgres guide: https://testcontainers.com/guides/testing-spring-boot-rest-api-using-testcontainers/",
        "Render multi-service deployment: https://render.com/docs/multi-service-architecture",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    doc.save(TECH_DOCX)
    return TECH_DOCX


if __name__ == "__main__":
    print(build_wireframes_doc())
    print(build_technical_doc())
