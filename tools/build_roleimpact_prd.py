from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parents[1] / "docs"
OUTPUT_PATH = OUTPUT_DIR / "RoleImpact_PRD_v0.1.docx"

# standard_business_brief preset, with a restrained product-blue override.
FONT = "Calibri"
NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
BODY = "202A33"
MUTED = "5E6B78"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E0E7"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "E8F3EC"
PALE_AMBER = "FFF4D6"
PALE_RED = "FBE9E7"
GREEN = "23643B"
AMBER = "8A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, *, size=None, bold=None, italic=None, color=BODY, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.color.rgb = rgb(color)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, *, color=MID_GRAY, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Column widths must sum to {CONTENT_WIDTH_DXA}, got {sum(widths_dxa)}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_table_rows_together(table):
    """Keep a short table on one page by chaining its paragraphs."""
    for row_index, row in enumerate(table.rows):
        keep_next = row_index < len(table.rows) - 1
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = keep_next


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_paragraph_border(paragraph, *, side="bottom", color=BLUE, size="10", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    element = OxmlElement(f"w:{side}")
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), space)
    element.set(qn("w:color"), color)
    p_bdr.append(element)


def shade_paragraph(paragraph, fill=PALE_BLUE, left_border=BLUE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), left_border)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_numbering_definition(doc, *, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "264")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(r_fonts)
        level.append(r_pr)

    abstract.append(level)
    # OOXML requires abstract numbering definitions to appear before concrete
    # numbering instances. Insert before the first w:num instead of appending.
    first_num_index = next(
        (idx for idx, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def add_list_item(doc, text: str, num_id: int, *, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.10
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, size=11, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=11)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11)
    return p


def add_body(doc, text: str, *, bold_lead: str | None = None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=11, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=11)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11)
    return p


def add_callout(doc, label: str, text: str, *, fill=PALE_BLUE, border=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.10
    shade_paragraph(p, fill=fill, left_border=border)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=11, bold=True, color=border)
    body_run = p.add_run(text)
    set_run_font(body_run, size=11, color=BODY)
    return p


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int], *, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = value
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=9.5, bold=True, color=NAVY)

    for row in table.rows[1:]:
        for idx, cell in enumerate(row.cells):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            for run in p.runs:
                set_run_font(run, size=9.2, color=BODY)
            if idx == 0 and len(headers) > 1:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = rgb(DARK_BLUE)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = doc.styles["Title"]
    title.font.name = FONT
    title._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title_p_pr = title._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = FONT
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    subtitle.font.size = Pt(13)
    subtitle.font.italic = False
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("ROLEIMPACT  |  PRODUCT REQUIREMENTS DOCUMENT")
        set_run_font(run, size=8.5, bold=True, color=MUTED)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        prefix = fp.add_run("Draft v0.1  |  Page ")
        set_run_font(prefix, size=9, color=MUTED)
        add_page_number(fp)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    return p


def add_metadata_line(doc, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, size=10, bold=True, color=NAVY)
    r2 = p.add_run(value)
    set_run_font(r2, size=10, color=BODY)
    return p


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    bullet_id = add_numbering_definition(doc, kind="bullet")

    core = doc.core_properties
    core.title = "RoleImpact Product Requirements Document"
    core.subject = "MVP product requirements for the RoleImpact access change impact simulator"
    core.author = "RoleImpact Project"
    core.keywords = "RoleImpact, PRD, access management, impact simulation, workflow continuity"
    core.comments = "Draft v0.1 for product review"

    # First-page memo masthead.
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(3)
    kr = kicker.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    set_run_font(kr, size=9, bold=True, color=BLUE)

    title = doc.add_paragraph("RoleImpact", style="Title")
    title.paragraph_format.keep_with_next = True
    subtitle = doc.add_paragraph(
        "Access Change Impact Simulator for Technical Access and Business Workflow Continuity",
        style="Subtitle",
    )
    subtitle.paragraph_format.keep_with_next = True

    add_metadata_line(doc, "Version", "Draft v0.1")
    add_metadata_line(doc, "Date", "August 12, 2026")
    add_metadata_line(doc, "Status", "For product review")
    add_metadata_line(doc, "Product stage", "Portfolio MVP")
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(4)
    rule.paragraph_format.space_after = Pt(10)
    set_paragraph_border(rule, color=BLUE, size="12", space="4")

    add_heading(doc, "1. Executive Summary", 1)
    add_callout(
        doc,
        "Product promise",
        "Before an employee, role, or permission is removed, RoleImpact shows what access disappears, which business workflows become blocked or fragile, why the impact occurs, and what minimal reassignment could restore coverage.",
    )
    add_body(
        doc,
        "Organizations often understand access changes only as lists of roles and permissions. Managers still need to translate those technical changes into operational questions: Can vendor payments still be approved? Is there another qualified release approver? Did the change create a single-person dependency? RoleImpact closes that explanation gap through safe, isolated simulations.",
    )
    add_body(
        doc,
        "The MVP is a decision-support application built around a synthetic organization. It will not connect to production identity systems or execute changes. Its purpose is to demonstrate full-stack product development, relationship-rich data modeling, deterministic impact analysis, graph exploration, explainable recommendations, and business-oriented UX.",
    )

    add_heading(doc, "1.1 MVP decision", 2)
    add_body(
        doc,
        "Proceed with the original RoleImpact direction, while making business workflow consequences a prominent result rather than limiting the product to a technical permission graph. Existing identity-governance products validate the problem; RoleImpact will use its own data model, interaction flow, scoring rules, visual design, and fictional scenarios.",
    )

    add_heading(doc, "2. Product Definition", 1)
    add_heading(doc, "2.1 Vision", 2)
    add_body(
        doc,
        "Make access-change decisions understandable before they become operational incidents. A reviewer should be able to move from a proposed change to a defensible decision without manually tracing roles, permissions, applications, employees, and workflows across separate systems.",
    )

    add_heading(doc, "2.2 Problem statement", 2)
    add_body(
        doc,
        "When an employee changes teams, leaves the company, loses a role, or has a permission removed, the direct technical effect may be clear while the business effect remains hidden. The same permission can support multiple applications and workflow steps, inherited roles can obscure the source of access, and a seemingly safe revocation can remove the last eligible approver for a critical process.",
    )

    add_heading(doc, "2.3 Product goals", 2)
    for item in (
        "Let an access manager simulate a proposed change without modifying the source organization.",
        "Calculate both direct technical impact and indirect workflow impact.",
        "Explain every important conclusion with a traceable relationship path.",
        "Identify blocked workflows, degraded coverage, and single-person dependencies.",
        "Offer one or two minimal, deterministic mitigation options when a safe option exists.",
        "Provide a polished end-to-end demonstration that can be understood by technical and non-technical interviewers.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "2.4 Non-goals for the MVP", 2)
    for item in (
        "Connecting to a real identity provider, HR system, cloud account, ERP, or production application.",
        "Provisioning, revoking, or approving real access.",
        "Replacing a complete IAM, IGA, PAM, or compliance platform.",
        "Machine-learning-based role mining, anomaly detection, or recommendation generation.",
        "Supporting arbitrary enterprise policy languages or every RBAC and ABAC edge case.",
        "Building multi-tenant billing, enterprise administration, or a marketplace of integrations.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "2.5 Product principles", 2)
    principles = [
        ["Explain before visualizing", "The graph supports the decision; it is not the decision. Results lead with plain-language consequences."],
        ["No hidden reasoning", "Every impact and recommendation includes the exact people, roles, permissions, rules, and workflow steps involved."],
        ["Simulation is isolated", "The baseline organization remains unchanged unless a future version explicitly implements approved changes."],
        ["Business consequence first", "A blocked approval or fragile deployment process is more important than a raw count of removed edges."],
        ["Small, credible scope", "Three workflows and a compact fictional organization are preferable to shallow support for dozens of systems."],
    ]
    add_table(doc, ["Principle", "Meaning"], principles, [2700, 6660])

    add_heading(doc, "3. Users and Jobs to Be Done", 1)
    add_heading(doc, "3.1 Primary persona", 2)
    add_callout(
        doc,
        "Primary user",
        "An IT or access manager who evaluates a requested access change before approving it. The user understands employees, roles, and applications but should not need to manually inspect every permission relationship.",
    )
    add_body(doc, "Primary job to be done: When I am reviewing an access removal, help me understand what will break and how to preserve necessary coverage so I can approve, reject, or revise the request confidently.")

    add_heading(doc, "3.2 Secondary persona", 2)
    add_body(
        doc,
        "A business or team manager who understands operational workflows but not the underlying access model. This user mainly consumes the business-impact explanation, reviews suggested coverage changes, and shares the resulting report.",
    )

    add_heading(doc, "3.3 Core user stories", 2)
    user_stories = [
        "As an access manager, I want to offboard an employee in a simulation so I can see all access and workflows that depend on that person.",
        "As an access manager, I want to revoke one role from one employee so I can isolate the impact of a specific request.",
        "As a role owner, I want to remove a permission from a role so I can see everyone and every workflow affected by the role definition change.",
        "As a reviewer, I want a before-and-after explanation so I can verify why a workflow became blocked or degraded.",
        "As a manager, I want to test a proposed mitigation so I can compare the original change with a safer alternative.",
        "As a reviewer, I want to save and export a scenario so I can document the decision and discuss it with others.",
    ]
    for story in user_stories:
        add_list_item(doc, story, bullet_id)

    add_heading(doc, "4. Exact User Flow", 1)
    add_callout(
        doc,
        "Primary path",
        "Dashboard > Create simulation > Select change and target > Review current access > Run simulation > Inspect impact > Test recommendation > Save or export.",
    )

    flow_steps = [
        ("Open dashboard", "Review organization health, critical workflows, single-person dependencies, and recent scenarios. Select New Simulation."),
        ("Choose change", "Select Offboard employee, Revoke role from employee, or Remove permission from role."),
        ("Select target", "Search for the employee, role, or permission. Inspect the target's current roles, applications, permissions, and workflow participation."),
        ("Confirm proposal", "Review the exact change and its direct baseline relationships. Confirm that no real access will be modified."),
        ("Run analysis", "Apply the proposed change to an isolated copy of the organization state and recalculate access, eligibility, constraints, and workflow coverage."),
        ("Review results", "Read the executive verdict, business impact, technical impact, explanation paths, and before/after graph."),
        ("Test mitigation", "Choose a recommended reassignment, branch the scenario, rerun the engine, and compare outcomes."),
        ("Preserve decision", "Name the scenario, add optional notes, save it to history, or export a review report."),
    ]
    flow_num_id = add_numbering_definition(doc, kind="decimal")
    for label, detail in flow_steps:
        p = add_list_item(doc, f"{label}: {detail}", flow_num_id)
        p.paragraph_format.keep_together = True

    add_heading(doc, "4.1 Secondary exploration flow", 2)
    add_body(
        doc,
        "A user may browse the Organization Explorer, open an employee, role, application, permission, or workflow detail page, and select Simulate a Change. The simulation form opens with that entity preselected. This keeps the explorer actionable while preserving the same analysis engine and result flow.",
    )

    add_heading(doc, "5. Information Architecture and Screens", 1)
    screens = [
        ["Dashboard", "Organization summary, health indicators, recent simulations", "Start a new simulation"],
        ["Organization Explorer", "Searchable relationship graph and filtered entity lists", "Open an entity or start a contextual simulation"],
        ["Entity Details", "Profile, effective access, relationships, supported workflows", "Inspect or simulate a change"],
        ["New Simulation", "Change type, target selectors, current-state preview", "Run impact analysis"],
        ["Simulation Results", "Verdict, business/technical impact, explanations, graph, recommendations", "Test a mitigation or save the scenario"],
        ["Scenario History", "Saved scenarios, dates, targets, verdicts, notes", "Reopen, duplicate, compare, or export"],
    ]
    add_table(doc, ["Screen", "Purpose", "Primary action"], screens, [2000, 4400, 2960])

    add_heading(doc, "5.1 Dashboard content", 2)
    for item in (
        "Organization counts: employees, teams, roles, applications, permissions, and workflows.",
        "Workflow health: operational, degraded, and blocked counts.",
        "Single-person dependencies: workflow steps with only one eligible actor.",
        "Recent simulations with change type, target, verdict, and timestamp.",
        "One prominent New Simulation action.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "5.2 Results-page hierarchy", 2)
    for item in (
        "Executive verdict: a plain-language summary such as 1 workflow blocked, 1 degraded, and 3 permissions removed.",
        "Business impact: affected workflow steps, coverage gaps, constraint failures, and operational consequences.",
        "Technical impact: removed roles, permissions, applications, resources, and affected identities.",
        "Explanation path: the exact relationship chain that produced each important result.",
        "Visual graph: affected nodes and removed edges, with a Before / After toggle.",
        "Recommendations: minimal safe options with rationale, trade-offs, and a Test This Recommendation action.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "6. Functional Requirements", 1)
    add_body(doc, "Priority definition: P0 requirements are required for the core demonstration. P1 requirements materially improve the product but may follow after the first complete vertical slice.")
    requirements = [
        ["FR-01", "Display organization counts, workflow health, weak coverage, and recent scenarios on the dashboard.", "P0"],
        ["FR-02", "Provide searchable and filterable lists for employees, teams, roles, permissions, applications, resources, and workflows.", "P0"],
        ["FR-03", "Display an entity detail page with direct relationships, effective access, and connected workflow steps.", "P0"],
        ["FR-04", "Allow a user to start a simulation from the dashboard or from an entity detail page.", "P1"],
        ["FR-05", "Support exactly three MVP change types: offboard employee, revoke role from employee, and remove permission from role.", "P0"],
        ["FR-06", "Show the selected target's current state and the exact proposed change before analysis begins.", "P0"],
        ["FR-07", "Run every proposed change against an isolated scenario state without modifying the baseline data.", "P0"],
        ["FR-08", "Calculate direct and inherited access before and after the proposed change.", "P0"],
        ["FR-09", "Recalculate eligible actors for every affected workflow step and evaluate workflow constraints.", "P0"],
        ["FR-10", "Classify each workflow as operational, degraded, or blocked and calculate an overall impact severity.", "P0"],
        ["FR-11", "Generate a plain-language explanation and relationship path for each blocked or degraded workflow.", "P0"],
        ["FR-12", "Display a focused before/after graph showing affected entities and removed or proposed relationships.", "P0"],
        ["FR-13", "Generate one or two deterministic mitigation options when an eligible candidate exists.", "P0"],
        ["FR-14", "Allow a recommendation to be tested as a branched simulation and compare the two outcomes.", "P1"],
        ["FR-15", "Save simulation inputs, outputs, timestamps, status, and optional reviewer notes.", "P0"],
        ["FR-16", "Reopen and duplicate saved simulations from scenario history.", "P1"],
        ["FR-17", "Export a readable impact report containing the proposal, verdict, evidence, and tested recommendation.", "P1"],
    ]
    add_table(doc, ["ID", "Requirement", "Priority"], requirements, [1000, 7160, 1200])

    add_heading(doc, "7. Impact Evaluation Rules", 1)
    add_heading(doc, "7.1 Workflow eligibility", 2)
    add_body(doc, "An employee is eligible for a workflow step only when all applicable rules are satisfied:")
    for item in (
        "The employee is active in the simulated state.",
        "The employee has the required capability through an effective permission and active role assignment.",
        "Any department, region, shift, or application prerequisites are satisfied.",
        "Assigning the employee to the step would not violate a separation-of-duties constraint.",
        "The employee is not explicitly excluded from the workflow or step.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "7.2 Workflow status", 2)
    status_rows = [
        ["Operational", "Every required step has the configured minimum number of eligible actors, and all constraints can be satisfied.", "Green"],
        ["Degraded", "The workflow remains executable, but at least one step falls below its resilience threshold or becomes a single-person dependency.", "Amber"],
        ["Blocked", "At least one required step has no eligible actor, or the workflow cannot satisfy a mandatory constraint.", "Red"],
    ]
    status_table = add_table(doc, ["Status", "Definition", "Display"], status_rows, [1700, 6260, 1400])
    keep_table_rows_together(status_table)
    set_cell_shading(status_table.rows[1].cells[0], PALE_GREEN)
    set_cell_shading(status_table.rows[2].cells[0], PALE_AMBER)
    set_cell_shading(status_table.rows[3].cells[0], PALE_RED)

    add_heading(doc, "7.3 Overall impact severity", 2)
    for item in (
        "Critical: a critical workflow becomes blocked.",
        "High: any non-critical workflow becomes blocked, or a critical workflow becomes degraded.",
        "Medium: a non-critical workflow becomes degraded or a new single-person dependency is introduced.",
        "Low: technical access changes occur without reducing workflow executability or resilience.",
        "None: the scenario produces no effective change.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "7.4 Evaluation sequence", 2)
    sequence = [
        "Load the immutable baseline organization state.",
        "Create a scenario copy and apply only the proposed change.",
        "Recompute effective roles and permissions for affected identities.",
        "Traverse from changed entities to connected applications, resources, capabilities, and workflow steps.",
        "Recompute eligible actors and evaluate coverage and separation-of-duties constraints.",
        "Compare baseline and scenario states to identify added, removed, blocked, and degraded relationships.",
        "Generate the verdict, explanation paths, graph diff, and eligible mitigation candidates.",
    ]
    evaluation_num_id = add_numbering_definition(doc, kind="decimal")
    for item in sequence:
        add_list_item(doc, item, evaluation_num_id)

    add_heading(doc, "7.5 Recommendation rules", 2)
    add_body(doc, "Recommendations are deterministic and limited to minimal access changes. A candidate must be active, satisfy organizational prerequisites, cover the affected shift or region, and introduce no separation-of-duties conflict. Candidates are ranked by:")
    for item in (
        "Fewest new permissions required.",
        "Existing membership in the appropriate department or team.",
        "Existing access to the relevant application.",
        "Ability to restore the greatest number of affected workflow steps.",
        "Improvement in backup coverage without creating excessive access.",
    ):
        add_list_item(doc, item, bullet_id)
    add_callout(doc, "Guardrail", "If no candidate satisfies the rules, RoleImpact must say that no safe automatic recommendation exists and identify the unmet constraints. It must not invent a recommendation.", fill=PALE_AMBER, border=AMBER)

    add_heading(doc, "8. Data Model", 1)
    entities = [
        ["Employee", "Person whose effective access and workflow eligibility can change", "Name, status, team, department, region, shift"],
        ["Team", "Organizational grouping used for ownership and candidate selection", "Name, department, manager"],
        ["Role", "Reusable access bundle assigned to employees", "Name, description, owner, sensitivity"],
        ["Permission", "Action granted through a role", "Action, application, resource, sensitivity"],
        ["Application", "System in which permissions operate", "Name, category, owner"],
        ["Resource", "Object governed by a permission", "Name, type, application"],
        ["Capability", "Business-facing ability mapped to one or more technical permissions", "Name, qualifying permissions"],
        ["Workflow", "Business process composed of ordered required steps", "Name, criticality, owner"],
        ["Workflow step", "Unit of work requiring a capability and coverage", "Order, capability, minimum actors, region/shift"],
        ["Constraint", "Rule restricting valid workflow assignments", "Type, participating steps, parameters"],
        ["Simulation", "Stored what-if scenario and its computed results", "Change, baseline version, verdict, evidence, notes"],
    ]
    add_table(doc, ["Entity", "Purpose", "Key fields"], entities, [1800, 3850, 3710])

    add_heading(doc, "8.1 Core relationships", 2)
    for item in (
        "Employee belongs to Team and may hold many Roles.",
        "Role grants many Permissions; the same Permission may appear in multiple Roles.",
        "Permission belongs to an Application and controls an action on a Resource.",
        "Capability is satisfied by one or more qualifying Permissions.",
        "Workflow contains ordered Workflow Steps; each step requires one Capability.",
        "Constraint connects workflow steps, employees, or organizational attributes.",
        "Simulation references a baseline state, applies a change set, and stores a result set without modifying the baseline.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "9. Fictional Demo Organization", 1)
    add_body(doc, "The MVP will ship with a seeded fictional company named Harborline Commerce. The dataset should be large enough to create realistic relationship chains but small enough to understand during a portfolio demonstration.")
    demo_rows = [
        ["Employees", "25", "Finance, Support, Engineering, Operations, and Security"],
        ["Teams", "5", "One team per major department"],
        ["Roles", "8", "Examples: Finance Analyst, Finance Approver, Support Agent, Release Approver"],
        ["Applications", "6", "Examples: LedgerPro, Vendor Portal, SupportDesk, DeployHub"],
        ["Permissions", "18-24", "Read, create, approve, refund, deploy, and administrative actions"],
        ["Workflows", "3", "Vendor Payment, Customer Refund, Production Deployment"],
        ["Workflow steps", "10-12", "Three to five steps per workflow"],
    ]
    add_table(doc, ["Object", "Target size", "Notes"], demo_rows, [1800, 1600, 5960])

    add_heading(doc, "9.1 Required workflow rules", 2)
    for item in (
        "Vendor Payment: the payment creator cannot approve the same payment; payments above $25,000 require an eligible Finance Approver; evening coverage requires at least one qualified approver.",
        "Customer Refund: refunds above $10,000 require a senior approver who did not create the refund request.",
        "Production Deployment: the release requester cannot approve the same deployment; production approval requires an active engineer or release manager with deploy approval capability.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "10. Demonstration Scenarios", 1)
    add_heading(doc, "10.1 Primary portfolio scenario", 2)
    add_callout(doc, "Proposed change", "Revoke the Finance Approver role from Priya Sharma.")
    scenario_steps = [
        "Before the change, Priya is the only evening-shift employee who can approve vendor payments above $25,000.",
        "The scenario removes the role and its effective payment approval permissions from Priya.",
        "RoleImpact reports that the Vendor Payment workflow is blocked during evening coverage and that Month-End Close has become a single-person dependency.",
        "The explanation traces Priya > Finance Approver > payment.approve > High-Value Payment Approval > Vendor Payment.",
        "RoleImpact recommends Bob Chen only if he satisfies the department, application, shift, and separation-of-duties rules.",
        "The reviewer tests the recommendation and compares the original blocked state with the restored state.",
    ]
    scenario_num_id = add_numbering_definition(doc, kind="decimal")
    for step in scenario_steps:
        add_list_item(doc, step, scenario_num_id)

    add_heading(doc, "10.2 Supporting scenarios", 2)
    supporting = [
        ["Offboard employee", "Offboard an engineering lead who is the only production release approver for one region.", "Production Deployment becomes degraded or blocked."],
        ["Remove permission from role", "Remove refund.approve from the Senior Support role.", "All employees holding the role are recalculated; high-value refund approval loses coverage."],
        ["No material impact", "Revoke a redundant viewer role while equivalent access remains inherited elsewhere.", "Technical graph changes, but workflows remain operational and severity is low."],
    ]
    add_table(doc, ["Change type", "Scenario", "Expected result"], supporting, [1900, 4300, 3160])

    add_heading(doc, "11. Non-Functional Requirements", 1)
    nfrs = [
        ["NFR-01", "Explainability", "Every blocked or degraded result must include a human-readable explanation and a traceable relationship path."],
        ["NFR-02", "Determinism", "The same baseline and scenario input must produce the same impact and recommendation output."],
        ["NFR-03", "Performance", "A simulation over the seeded MVP dataset should return in under two seconds on a typical development machine."],
        ["NFR-04", "Isolation", "Running, duplicating, or testing a simulation must not mutate the baseline organization data."],
        ["NFR-05", "Usability", "A first-time reviewer should complete the primary scenario without documentation in under three minutes."],
        ["NFR-06", "Accessibility", "Core interactions must support keyboard navigation, visible focus, sufficient contrast, and text alternatives for graph conclusions."],
        ["NFR-07", "Responsive layout", "The product must remain usable on common laptop widths; mobile optimization is not required for MVP."],
        ["NFR-08", "Auditability", "Saved scenarios must preserve the baseline version, proposed change, result, recommendation test, timestamp, and reviewer notes."],
        ["NFR-09", "Error handling", "Invalid or contradictory data must produce a clear error or inconclusive result rather than a fabricated conclusion."],
    ]
    add_table(doc, ["ID", "Quality", "Requirement"], nfrs, [1050, 1750, 6560])

    add_heading(doc, "12. MVP Acceptance Criteria", 1)
    add_heading(doc, "12.1 End-to-end acceptance scenario", 2)
    acceptance = [
        "Given the seeded Harborline Commerce dataset, the user can open the dashboard and start a new simulation.",
        "The user can select Revoke role from employee, choose Priya Sharma, and select Finance Approver.",
        "The preview accurately displays Priya's current roles, affected permissions, applications, and connected workflows.",
        "Running the simulation leaves baseline data unchanged and returns a severity and workflow status summary.",
        "The results identify the affected workflow step, technical access removed, and the complete explanation path.",
        "The graph differentiates unchanged, removed, blocked, and degraded relationships and supports before/after inspection.",
        "When a safe candidate exists, RoleImpact provides a deterministic recommendation with a rationale.",
        "Testing the recommendation creates a branch and displays a before/after comparison without overwriting the original scenario.",
        "The user can save the scenario and reopen it from Scenario History with all inputs and results intact.",
    ]
    for item in acceptance:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "12.2 Definition of done", 2)
    for item in (
        "All P0 functional requirements are implemented and demonstrable.",
        "The three required change types produce correct baseline-versus-scenario diffs.",
        "Automated tests cover impact classification, workflow eligibility, separation of duties, and recommendation ranking.",
        "The primary scenario can be completed without manual data changes or developer intervention.",
        "Empty, loading, error, and no-recommendation states are intentionally designed.",
        "The seeded data, screenshots, README, architecture explanation, and demo instructions are ready for portfolio review.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "13. Success Metrics", 1)
    success_rows = [
        ["Task completion", "A new reviewer completes the primary scenario in under 3 minutes without guidance.", "Informal usability test with 3-5 people"],
        ["Explanation clarity", "At least 80% of reviewers correctly explain why the workflow was affected after reading the result.", "One comprehension question after the demo"],
        ["Engine correctness", "All authored rule-engine test cases produce the expected workflow status and explanation path.", "Automated test suite"],
        ["Responsiveness", "Seeded simulations return in under 2 seconds on the development machine.", "Performance test"],
        ["Portfolio readiness", "A visitor can understand the problem, run the demo, and inspect implementation decisions from the repository.", "README and demo review checklist"],
    ]
    add_table(doc, ["Metric", "Target", "Measurement"], success_rows, [1900, 4400, 3060])

    add_heading(doc, "14. Delivery Plan", 1)
    delivery = [
        ["Phase 1: Product foundation", "Approve PRD; define screen wireframes, data model, business rules, and technical architecture.", "Approved product and implementation plan"],
        ["Phase 2: Vertical slice", "Seed one workflow; implement role revocation; compute impact; display a basic result.", "One complete simulation path"],
        ["Phase 3: MVP engine", "Add all three change types, three workflows, constraints, severity, explanations, and recommendations.", "Correct deterministic analysis"],
        ["Phase 4: Product experience", "Add explorer, graph diff, scenario history, comparison, loading/error states, and polished UI.", "Portfolio-quality application"],
        ["Phase 5: Verification", "Complete automated tests, usability checks, documentation, deployment, and demo walkthrough.", "Publicly reviewable portfolio project"],
    ]
    add_table(doc, ["Phase", "Scope", "Exit condition"], delivery, [2100, 4650, 2610])

    post_mvp_heading = add_heading(doc, "14.1 Post-MVP candidates", 2)
    post_mvp_heading.paragraph_format.page_break_before = True
    for item in (
        "CSV import for employees, role assignments, and permissions.",
        "Additional scenario types such as transfer employee, modify team membership, or add a new role assignment.",
        "Side-by-side comparison of multiple mitigation strategies.",
        "Time-bound changes and future-dated simulations.",
        "Optional integration with a sample identity-provider API in read-only mode.",
        "Natural-language summaries generated from deterministic evidence, with the evidence remaining authoritative.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "15. Risks and Mitigations", 1)
    risks = [
        ["Scope expands into a full IAM platform", "Time is spent on authentication, provisioning, and integrations instead of impact analysis.", "Keep the MVP synthetic and read-only; enforce the non-goals."],
        ["Graph becomes decorative or unreadable", "A large graph impresses visually but does not help the decision.", "Lead with summaries and evidence paths; default the graph to affected nodes only."],
        ["Recommendation logic appears arbitrary", "Users cannot trust why a candidate was chosen.", "Use deterministic ranking and display every eligibility and exclusion reason."],
        ["Workflow model becomes too complex", "The project stalls on enterprise-grade process modeling.", "Support three explicit workflows and a small set of constraints."],
        ["Synthetic data feels unrealistic", "The demo looks like a toy even if the implementation is strong.", "Use coherent names, teams, shifts, role overlap, inherited access, and both positive and negative scenarios."],
        ["Technical stack drives the product", "Architecture decisions are made before the domain model and user experience are stable.", "Approve this PRD and wireframes before finalizing implementation choices."],
    ]
    add_table(doc, ["Risk", "Consequence", "Mitigation"], risks, [2400, 3270, 3690])

    add_heading(doc, "16. Decisions and Review Questions", 1)
    add_heading(doc, "16.1 Decisions currently treated as locked", 2)
    for item in (
        "The product name is RoleImpact for the MVP.",
        "The primary persona is an IT or access manager reviewing a proposed access removal.",
        "The MVP is a safe simulator and does not execute real changes.",
        "The MVP supports three change types and three business workflows.",
        "Business workflow impact and technical access impact are equally visible, with business consequences shown first.",
        "The analysis and recommendation engines are deterministic and explainable.",
        "The product uses a seeded fictional organization rather than live integrations.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "16.2 Items to confirm during PRD review", 2)
    review_questions = [
        "Should Test This Recommendation remain P1, or is it important enough to be P0 for the first public demo?",
        "Should report export be included in the MVP or deferred until the core web experience is complete?",
        "Are Vendor Payment, Customer Refund, and Production Deployment the right three workflows for the fictional company?",
        "Should the dashboard show pre-existing organizational risks, or focus only on simulation history and organization counts?",
        "Is Harborline Commerce an acceptable name and domain for the seeded organization?",
    ]
    for item in review_questions:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "17. Glossary", 1)
    glossary = [
        ["Baseline", "The immutable current organization state against which a simulation is compared."],
        ["Capability", "A business-facing ability, such as Approve high-value payment, satisfied by one or more permissions."],
        ["Direct impact", "An entity or relationship explicitly removed by the proposed change."],
        ["Effective access", "Permissions an employee receives after role assignments and inheritance are resolved."],
        ["Explanation path", "The traceable chain connecting a changed entity to an operational result."],
        ["Indirect impact", "A downstream consequence discovered by traversing relationships from a direct change."],
        ["Separation of duties", "A rule preventing one actor from performing conflicting workflow steps."],
        ["Single-person dependency", "A workflow step for which only one eligible actor remains."],
        ["Workflow resilience", "The ability of a workflow to remain executable when one or more actors or permissions are unavailable."],
    ]
    add_table(doc, ["Term", "Definition"], glossary, [2500, 6860])

    add_callout(
        doc,
        "Next step after approval",
        "Convert the approved PRD into low-fidelity wireframes and a technical design covering the final data schema, impact-engine contract, API boundaries, test strategy, and implementation stack.",
        fill=PALE_GREEN,
        border=GREEN,
    )

    add_header_footer(doc)

    # Preserve headings with following content and discourage row splitting.
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path.resolve())
